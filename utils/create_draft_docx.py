import docx
import os

def create_draft():
    doc = docx.Document()
    
    # Title
    doc.add_heading('Multi-Agent Large Language Models for Objective Theory Evaluation in Neuroscience', 0)
    
    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'Neuroscience theories, ranging from neurophysiological principles like Hierarchical Predictive Coding '
        '(HPC) to biophysical pathologies like Schizophrenia (ScZ) circuit deficits, are supported by vast, '
        'heterogeneous, and often conflicting literature. Traditional meta-analyses struggle with selection bias '
        'and the dimensionality of tracking dozens of mechanistic factors across hundreds of papers. Here, we '
        'introduce a novel, multi-agent AI framework for the objective evaluation of scientific theories. By '
        'defining strict, domain-specific glossaries of theoretical factors, we deploy an ensemble of Large '
        'Language Models (LLMs)—including DeepSeek-R1, Qwen 3.5, and Gemini 1.5 Pro—to autonomously read, score, '
        'and evaluate the literature. We demonstrate the efficacy of this "Algorithmic Knowledge Synthesis" across '
        'multiple neuroscientific domains: (1) mapping the empirical consensus of HPC, (2) evaluating competing '
        'biophysical models of ScZ circuit deficits, and (3) a third domain integrating the physics of the brain. '
        'By utilizing cross-agent agreement as a "Certainty Index," we provide a high-resolution, unbiased map of '
        'the current scientific consensus and the primary frontiers of debate.'
    )
    
    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_heading('1. The Fragmentation of Neuroscience Theories', level=2)
    doc.add_paragraph(
        'Modern neuroscience is characterized by the integration of computational, biological, and physical sciences. '
        'Grand theories, such as Hierarchical Predictive Coding or the Excitation/Inhibition (E/I) imbalance hypothesis '
        'in Schizophrenia, rely on evidence spanning single-unit electrophysiology, macroscopic neuroimaging (MEG/EEG), '
        'and in silico biophysical simulations. Consequently, the literature evaluating these theories is massive and '
        'heavily fragmented. Human-led literature reviews are increasingly bottlenecked by dimensionality limits—the '
        'inability to simultaneously track 30+ interacting factors across hundreds of papers—and inherent selection bias.'
    )
    
    doc.add_heading('2. The Solution: Multi-Agent LLM Evaluation', level=2)
    doc.add_paragraph(
        'To bridge this gap, we present a standardized framework for the algorithmic evaluation of scientific literature. '
        'We construct explicit "Factor Glossaries" (e.g., the 36-factor TcGLO glossary for Predictive Coding) that define '
        'the structural, functional, and methodological pillars of a given theory. We then provide these glossaries as '
        'specialized skills to an ensemble of LLMs. By comparing the evaluations across models with different architectures '
        'and training paradigms (DeepSeek, Qwen, Gemini), we establish a robust "Consensus Score" and measure "Agent Variance" '
        'to quantify the certainty of the scientific community regarding specific mechanistic claims.'
    )

    doc.add_heading('3. Case Study 1: Hierarchical Predictive Coding (Neurophysiology)', level=2)
    doc.add_paragraph(
        'Our first application focuses on the neurophysiological dynamics of HPC. We evaluate over 200 empirical and '
        'computational studies against a 36-factor glossary, identifying universal motifs such as feedforward Gamma '
        'error propagation and alpha/beta-mediated predictive suppression.'
    )

    doc.add_heading('4. Case Study 2: Schizophrenia Circuit Deficits (Biophysics & Pathology)', level=2)
    doc.add_paragraph(
        'Our second application transitions to pathology. We define a new glossary evaluating the cellular basis of ScZ, '
        'specifically focusing on the competing and complementary theories of Parvalbumin (PV) deficit-driven Gamma reduction '
        'versus Somatostatin (SST) driven Beta enhancement, evaluating literature connecting microcircuit anatomy to macroscopic symptoms.'
    )

    doc.add_heading('5. Case Study 3: [To Be Determined - Physics/Connectivity]', level=2)
    doc.add_paragraph(
        '[Placeholder for the third use case, potentially focusing on connectomics, criticality, or the physics of brain networks].'
    )

    doc.add_heading('Methods', level=1)
    doc.add_heading('1. Glossary Definition and Prompt Engineering', level=2)
    doc.add_paragraph(
        'For each theory, a rigorous glossary of factors is defined, categorizing them into qualitative, quantitative, and '
        'methodological tags. These are injected into the context window of the LLM ensemble along with strict scoring rubrics.'
    )

    doc.add_heading('2. The Agent Ensemble', level=2)
    doc.add_paragraph(
        'We utilize a mixture of open-weights and proprietary models to ensure evaluation diversity.\n'
        '- DeepSeek-R1: Emphasizing deep reasoning traces.\n'
        '- Qwen 3.5 (122B): Providing dense parameter understanding.\n'
        '- Gemini 1.5 Pro: Offering massive context windows for full-text comprehension.'
    )

    doc.add_heading('3. Consensus and Certainty Metrics', level=2)
    doc.add_paragraph(
        'We define the Consensus Score as the mean evaluation across agents, and the Certainty Index as the inverse variance '
        'between agents. High consensus and high certainty indicate "settled science," while high variance indicates either '
        'literature ambiguity or algorithmic limitation.'
    )

    # Save to the Google Drive Sync Folder
    save_dir = "/Users/hamednejat/Google Drive/My Drive/Workspace/drafts"
    os.makedirs(save_dir, exist_ok=True)
    save_path = os.path.join(save_dir, 'Multi_Agent_Neuro_Eval.docx')
    doc.save(save_path)
    print(f"Draft saved to {save_path}")

if __name__ == "__main__":
    create_draft()
